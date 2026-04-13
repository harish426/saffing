import json
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from app.services.ai_services import GeminiService



class Resume:
    def __init__(self, json_data):
        # Handle case where resume data is wrapped in "data" key (e.g. from API response)
        if "data" in json_data and "personal_info" in json_data["data"]:
            self.data = json_data["data"]
        else:
            self.data = json_data
        self.doc = Document()
        self._setup_document()
        self.ai = GeminiService()

    def _setup_document(self):
        # --- Page Margins ---
        section = self.doc.sections[0]
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

        # --- Global Font ---
        style = self.doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(10)

    def set_font(self, run, font_name="Arial", size=11, bold=False, italic=False, color=None):
        run.font.name = font_name
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color

    def add_section_heading(self, text):
        p = self.doc.add_paragraph()
        run = p.add_run(text.upper())
        self.set_font(run, size=12, bold=True, color=RGBColor(0, 0, 139)) # Dark Blue
        run.font.underline = True # Underline only the text
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.keep_with_next = True

    def add_personal_info(self):
        info = self.data.get("personal_info", {})
        name = info.get("name", "")
        title = info.get("title", "")
        email = info.get("email", "")
        phone = info.get("phone", "")
        linkedin = info.get("linkedin", "")

        # Name
        if name:
            p_name = self.doc.add_paragraph()
            p_name.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run_name = p_name.add_run(name)
            self.set_font(run_name, size=22, bold=True)
        
        # Title
        if title:
            p_title = self.doc.add_paragraph()
            p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run_title = p_title.add_run(title)
            self.set_font(run_title, size=12, color=RGBColor(80, 80, 80)) # Dark Gray

        # Contact
        if phone or email or linkedin:
            p_contact = self.doc.add_paragraph()
            p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_text = " | ".join(p for p in [phone, email, linkedin] if p)
            run_contact = p_contact.add_run(contact_text)
            self.set_font(run_contact, size=10)

        

    def add_summary(self, job_description):
        summary = self.data.get("professional_summary", {})
        if not summary.get("overview") and not summary.get("key_highlights"):
            return

        try:
            # Fix: Use get_Ai_Subject for the overview text (assuming that's the intention for single-string tailoring)
            # and DO NOT overwrite the summary dict with the result string.
            tailored_overview = self.ai.get_Ai_Subject(job_description, summary.get("overview", ""))
            summary["overview"] = tailored_overview
        except Exception as e:
            # print(f"Error tailoring summary overview: {e}")
            pass
        self.add_section_heading("Professional Summary")
        
        overview = summary.get("overview", "")
        if overview:
            p = self.doc.add_paragraph(overview)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(8)
            if summary.get("key_highlights"):
                p.paragraph_format.keep_with_next = True

        highlights = summary.get("key_highlights", [])
        if highlights:
            try:    
                highlights = self.ai.generate_tailored_resume_content(highlights, job_description)
            except Exception as e:
                # print(e)
                highlights = summary.get("key_highlights", [])
            for i, item in enumerate(highlights):
                p = self.doc.add_paragraph(item, style='List Bullet')
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                # Glue the first bullet to the next paragraph to ensure Heading + Bullet 1 + Bullet 2 stay together
                if i == 0 and len(highlights) > 1:
                     p.paragraph_format.keep_with_next = True
            

    def add_education(self):
        edu = self.data.get("education", {})
        if edu and edu.get("degree")!="":
            self.add_section_heading("Education")
            degree = edu.get("degree", "")
            inst = edu.get("institution", "")
            loc = edu.get("location", "")
            
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(f"{degree} | {inst}, {loc}")
            self.set_font(run, bold=True)

    def add_skills(self, requirements):
        tech = self.data.get("technologies", {})
        if tech:
            try:
                tech = self.ai.tailor_technical_skills(tech, requirements)
            except Exception as e:
                tech = self.data.get("technologies", {})


            self.add_section_heading("Technical Skills")
            
            table = self.doc.add_table(rows=0, cols=2)
            table.style = 'Table Grid'
            table.autofit = True
            
            for category, skills in tech.items():
                row_cells = table.add_row().cells
                
                # Category formatting (Title Case, Bold)
                cat_name = category.replace("_", " ").title()
                run_cat = row_cells[0].paragraphs[0].add_run(cat_name)
                self.set_font(run_cat, bold=True, size=10)
                
                # Skills formatting
                if isinstance(skills, list):
                    skills_str = ", ".join(skills)
                else:
                    skills_str = str(skills)
                    
                run_skills = row_cells[1].paragraphs[0].add_run(skills_str)
                self.set_font(run_skills, size=10)

    def _add_job_entry(self, job):
        if job.get("client", "") == "" and job.get("role", "") == "":
            return
        client = job.get("client", "")
        role = job.get("role", "")
        duration = job.get("duration", "")
        location = job.get("location", "")
        
        # Line 1: Client, Location
        p_header = self.doc.add_paragraph()
        p_header.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY 
        
        run_client = p_header.add_run(client)
        self.set_font(run_client, bold=True, size=11)
        p_header.paragraph_format.keep_with_next = True
        
        if location:
            run_loc = p_header.add_run(f", {location}")
            self.set_font(run_loc, size=11, italic=True)

        # Line 2: Role (Left) ... Duration (Right)
        p_sub = self.doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_sub.paragraph_format.keep_with_next = True
        
        # Calculate usable width: 8.5 (Letter) - 0.75 (Left) - 0.5 (Right) = 7.25 inches
        tab_stops = p_sub.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(7.25), WD_TAB_ALIGNMENT.RIGHT)
        
        run_role = p_sub.add_run(role)
        self.set_font(run_role, bold=True, size=11)
        
        if duration:
            # Add Tab character + Duration Text (No parens)
            run_dur = p_sub.add_run(f"\t{duration}")
            self.set_font(run_dur, size=11)

        # Responsibilities
        resps = job.get("responsibilities", [])
        for r in resps:
            p_r = self.doc.add_paragraph(r, style='List Bullet')
            p_r.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_r.paragraph_format.space_after = Pt(2)
        
        # Spacer between jobs
        self.doc.add_paragraph() 

    def add_experience_part_one(self, job_description):
        experience = self.data.get("professional_experience", [])
        if not experience:
            return
        
        self.add_section_heading("Professional Experience")
        
        # First 2 experiences
        for job in experience[:2]:
            try:
                job_resposibilities = self.ai.generate_tailored_resume_content(job["responsibilities"], job_description)
                job["responsibilities"] = job_resposibilities
            except Exception:
                pass
            self._add_job_entry(job)

    def add_experience_part_two(self, job_description):
        experience = self.data.get("professional_experience", [])
        if experience and len(experience) > 2:
            # Remaining experiences
            for job in experience[2:]:
                try:
                    job_resposibilities = self.ai.generate_tailored_resume_content(job["responsibilities"], job_description)
                    job["responsibilities"] = job_resposibilities
                except Exception:
                    pass
                self._add_job_entry(job) 

    def add_projects(self):
        if self.data.get("projects",[]) == []:
            return
        projects =self.data.get("projects",[])
        if projects:
            self.add_section_heading("Projects")
            for project in projects:
                self._add_project_entry(project)
    def _add_project_entry(self, project):
        name = project.get("name", "")
        description = project.get("description", [])
        
        # Line 1: Project Name (Bold)
        p_header = self.doc.add_paragraph()
        p_header.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY 
        
        run_name = p_header.add_run(name)
        self.set_font(run_name, bold=True, size=11)
        p_header.paragraph_format.keep_with_next = True
        
        # Responsibilities
        for r in description:
            p_r = self.doc.add_paragraph(r, style='List Bullet')
            p_r.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_r.paragraph_format.space_after = Pt(2)
        
        # Spacer between projects
        self.doc.add_paragraph() 
    
    def add_certifications(self):
        if self.data.get("certifications", []) == [] or self.data.get("certifications", []) == {}:
            return
        certifications = self.data.get("certifications", [])
        if certifications:
            self.add_section_heading("Certifications")
            self._add_certification_entry(certifications)
    def _add_certification_entry(self, certifications):
        cert_names=[]
        
        for cert in certifications:
            cert_names.append(cert.get("name", "") + " (" + cert.get("institution", "") + ")")
        cert_string = ", ".join(cert_names)
        p = self.doc.add_paragraph(cert_string)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.set_font(p.runs[0] if p.runs else p.add_run(cert_string), size=10) 
    def save(self, output_path):
        self.doc.save(output_path)

        
       


def create_resume_from_json(json_path, output_path, job_description, requirements):
    if not os.path.exists(json_path):
        print(f"Error: JSON file not found at {json_path}")
        return

    with open(json_path, 'r') as f:
        data = json.load(f)
    resume = Resume(data)
    resume.add_personal_info()
    resume.add_summary(job_description)
    resume.add_education()
    resume.add_skills(requirements)
    resume.add_experience_part_one(job_description)
    resume.add_experience_part_two(job_description)
    resume.add_projects()
    resume.add_certifications()
    resume.save(output_path)
    

def generate_resume_buffer(resume_data, job_description, requirements):
    import io
    
    if isinstance(resume_data, str):
        # Check if it's a file path
        if os.path.exists(resume_data):
            with open(resume_data, 'r') as f:
                data = json.load(f)
        else:
             # Assume it's a JSON string
             try:
                 data = json.loads(resume_data)
             except json.JSONDecodeError:
                 # It was neither a valid path nor valid JSON
                 import logging
                 logger = logging.getLogger(__name__)
                 logger.error(f"Error: resume_data is neither a valid file path nor valid JSON string.")
                 return None
    else:
        # It's already a dict (from DB)
        data = resume_data

    resume = Resume(data)
    resume.add_personal_info()
    resume.add_summary(job_description)
    resume.add_education()
    resume.add_skills(requirements)
    resume.add_experience_part_one(job_description)
    resume.add_experience_part_two(job_description)
    resume.add_projects()
    resume.add_certifications()
    buffer = io.BytesIO()
    resume.save(buffer)
    buffer.seek(0)
    return buffer

if __name__ == "__main__":
    base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    json_source = os.path.join(base_dir, "data", "resume", "resume.json")
    output_tgt = os.path.join(base_dir, "Harish_Jamallamudi_Resume_Layout_v17.docx")
    
    sample_jd = """
    Looking for a GenAI Engineer with experience in RAG, Vector Databases (FAISS/Chroma), and LLM orchestration (LangChain). 
    Must be proficient in Python and SQL. Experience with cloud deployment is a plus.
    """
    requirements = ["Python", "SQL", "RAG", "AWS"]
    
    create_resume_from_json(json_source, output_tgt, sample_jd, requirements)
