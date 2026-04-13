import json
import os
import io
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

class FresherResume:
    def __init__(self, json_data):
        if "data" in json_data and "personal_info" in json_data["data"]:
            self.data = json_data["data"]
        else:
            self.data = json_data
        self.doc = Document()
        self._setup_document()

    def _setup_document(self):
        # --- Slim Margins ---
        section = self.doc.sections[0]
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.6)
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)

        # --- Base Font (Reverted to 8.5pt) ---
        style = self.doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(8.5)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def set_font(self, run, size=8.5, bold=False, italic=False, color=None):
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color

    def add_section_heading(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(text.upper())
        self.set_font(run, size=10, bold=True, color=RGBColor(0, 0, 0))
        
        # Border
        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '2')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        pBdr.append(bottom)
        pPr.append(pBdr)

    def add_hyperlink(self, paragraph, text, url, size=8.5, bold=False):
        # This gets access to the document relationships
        part = paragraph.part
        r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

        # Create the w:hyperlink tag
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)

        # Create a w:r element
        new_run = OxmlElement('w:r')

        # Create a w:rPr element
        rPr = OxmlElement('w:rPr')

        # Add styling to the run
        rStyle = OxmlElement('w:rStyle')
        rStyle.set(qn('w:val'), 'Hyperlink')
        rPr.append(rStyle)
        
        # Color blue and underline
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0000FF')
        rPr.append(color)
        
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
        
        # Set size and bold
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(size * 2))
        rPr.append(sz)
        
        if bold:
            b = OxmlElement('w:b')
            rPr.append(b)

        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)

        paragraph._p.append(hyperlink)
        return hyperlink

    def add_header(self):
        info = self.data.get("personal_info", {})
        name = info.get("name", "YOUR NAME")
        
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_name = p.add_run(name)
        self.set_font(run_name, size=16, bold=True)
        
        # Contact Info Line
        p_contact = self.doc.add_paragraph()
        p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Email & Phone (Normal Text)
        email = info.get("email", "")
        phone = info.get("phone", "")
        
        parts = []
        if email: parts.append(email)
        if phone: parts.append(phone)
        
        contact_text = " | ".join(parts)
        if contact_text:
            run_c = p_contact.add_run(contact_text)
            self.set_font(run_c, size=9) # Name/Links usually slightly different size
            if any([info.get("linkedin"), info.get("github"), info.get("portfolio")]):
                p_contact.add_run(" | ")

        # Helper to add links with separator
        def append_link(p, text, url, is_last=False):
            self.add_hyperlink(p, text, url, size=9)
            if not is_last:
                p.add_run(" | ")

        links_to_add = []
        if info.get("linkedin"): links_to_add.append(("LinkedIn", info.get("linkedin")))
        if info.get("github"): links_to_add.append(("Github", info.get("github")))
        if info.get("portfolio"): links_to_add.append(("Portfolio", info.get("portfolio")))

        for i, (text, url) in enumerate(links_to_add):
            append_link(p_contact, text, url, is_last=(i == len(links_to_add) - 1))

    def add_skills(self):
        tech = self.data.get("technologies", {})
        if tech:
            self.add_section_heading("Technical Skills")
            for category, skills in tech.items():
                p = self.doc.add_paragraph(style='List Bullet')
                p.paragraph_format.left_indent = Inches(0.2)
                run_cat = p.add_run(f"{category}: ")
                self.set_font(run_cat, bold=True)
                
                if isinstance(skills, list):
                    skills_str = ", ".join(skills)
                else:
                    skills_str = str(skills)
                p.add_run(skills_str)

    def add_experience(self):
        exp = self.data.get("professional_experience", [])
        if exp:
            self.add_section_heading("Professional Experience")
            for job in exp:
                p = self.doc.add_paragraph()
                p.paragraph_format.space_before = Pt(3)
                
                # Tab stop at right margin: 8.5" - 0.7" - 0.6" = 7.2"
                tab_stops = p.paragraph_format.tab_stops
                tab_stops.add_tab_stop(Inches(7.2), WD_TAB_ALIGNMENT.RIGHT)

                # Employer, Location | Role
                run_client = p.add_run(f"{job.get('client', '')}, {job.get('location', '')}")
                self.set_font(run_client, size=9, bold=True)
                run_sep = p.add_run(" | ")
                self.set_font(run_sep, size=9)
                run_role = p.add_run(job.get("role", ""))
                self.set_font(run_role, size=9, bold=False)
                
                # Add Tab and Duration on same line
                run_dur = p.add_run(f"\t{job.get('duration', '')}")
                self.set_font(run_dur, size=9)
                
                resps = job.get("responsibilities", [])
                for r in resps:
                    p_r = self.doc.add_paragraph(r, style='List Bullet')
                    p_r.paragraph_format.left_indent = Inches(0.2)

    def add_projects(self):
        projects = self.data.get("projects", [])
        if projects:
            self.add_section_heading("Projects")
            for proj in projects:
                p = self.doc.add_paragraph()
                p.paragraph_format.space_before = Pt(3)
                
                name = proj.get("name", "")
                url = proj.get("url")
                
                if url and "(" in name and ")" in name:
                    start_idx = name.find("(")
                    end_idx = name.find(")") + 1
                    
                    # Normal text before link
                    before = name[:start_idx]
                    if before:
                        run_before = p.add_run(before)
                        self.set_font(run_before, bold=True)
                    
                    # Hyperlink part
                    link_text = name[start_idx:end_idx]
                    self.add_hyperlink(p, link_text, url, bold=True)
                    
                    # Normal text after link
                    after = name[end_idx:]
                    if after:
                        run_after = p.add_run(after)
                        self.set_font(run_after, bold=True)
                elif url:
                    self.add_hyperlink(p, name, url, bold=True)
                else:
                    run_name = p.add_run(name)
                    self.set_font(run_name, bold=True)
                
                desc = proj.get("description", [])
                for d in desc:
                    p_d = self.doc.add_paragraph(d, style='List Bullet')
                    p_d.paragraph_format.left_indent = Inches(0.2)

    def add_education_custom(self):
        edu_list = self.data.get("education_list", [])
        if not edu_list:
            # Fallback for old format
            edu = self.data.get("education", {})
            if edu:
                edu_list = [edu]
        
        if edu_list:
            self.add_section_heading("Education")
            for edu in edu_list:
                p_inst = self.doc.add_paragraph()
                p_inst.paragraph_format.space_before = Pt(3)
                
                # Tab stop at right margin: 7.2"
                tab_stops = p_inst.paragraph_format.tab_stops
                tab_stops.add_tab_stop(Inches(7.2), WD_TAB_ALIGNMENT.RIGHT)

                # Institution (Bold)
                run_inst = p_inst.add_run(edu.get('institution', ''))
                self.set_font(run_inst, size=9, bold=True)
                
                # Location (Right Aligned, Bold)
                run_loc = p_inst.add_run(f"\t{edu.get('location', '')}")
                self.set_font(run_loc, size=9, bold=True)
                
                p_deg = self.doc.add_paragraph()
                # Tab stop at right margin: 7.2"
                tab_stops_deg = p_deg.paragraph_format.tab_stops
                tab_stops_deg.add_tab_stop(Inches(7.2), WD_TAB_ALIGNMENT.RIGHT)

                # Degree and Duration (Right Aligned)
                run_deg = p_deg.add_run(edu.get('degree', ''))
                self.set_font(run_deg, size=9, bold=False)
                
                run_dur_edu = p_deg.add_run(f"\t{edu.get('duration', '')}")
                self.set_font(run_dur_edu, size=9)

    def save(self, output_path):
        self.doc.save(output_path)

def generate_fresher_resume(data, output_path):
    resume = FresherResume(data)
    resume.add_header()
    resume.add_skills()
    resume.add_experience()
    resume.add_projects()
    resume.add_education_custom()
    resume.save(output_path)
    return output_path

def generate_fresher_resume_buffer(data):
    """Generates the resume and returns it as a BytesIO buffer."""
    resume = FresherResume(data)
    resume.add_header()
    resume.add_skills()
    resume.add_experience()
    resume.add_projects()
    resume.add_education_custom()
    
    buffer = io.BytesIO()
    resume.doc.save(buffer)
    buffer.seek(0)
    return buffer

if __name__ == "__main__":
    # Test block
    sample_json = os.path.join(os.path.dirname(__file__), "..", "..", "data", "resume", "resume.json")
    if os.path.exists(sample_json):
        with open(sample_json, 'r') as f:
            data = json.load(f)
        output = os.path.join(os.path.dirname(__file__), "..", "..", "fresher_resume_output.docx")
        generate_fresher_resume(data, output)
        print(f"Resume generated at: {output}")
